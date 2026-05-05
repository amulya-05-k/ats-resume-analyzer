import os
import re


def parse_resume(filepath: str) -> str:
    """Parse a resume file (PDF or DOCX) and return extracted text."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return _parse_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        return _parse_docx(filepath)
    else:
        raise ValueError(f'Unsupported file type: {ext}')


def _parse_pdf(filepath: str) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return _clean_text('\n'.join(text_parts))


def _parse_docx(filepath: str) -> str:
    from docx import Document

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    return _clean_text('\n'.join(paragraphs))


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    # Remove non-printable chars except newlines and tabs
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]', ' ', text)
    # Collapse multiple spaces/tabs
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse more than 2 consecutive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
