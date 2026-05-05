from io import BytesIO
import hashlib
import logging
from pathlib import Path
import re
from typing import Tuple

import pdfplumber
from docx import Document
from fastapi import UploadFile, HTTPException

from app.core.config import settings

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_MIME_TYPES_BY_EXTENSION = {
    ".pdf": {"application/pdf", "application/x-pdf", "application/acrobat"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    },
}
MAX_UPLOAD_BYTES = settings.max_upload_mb * 1024 * 1024


def _is_browser_default_mime_type(content_type: str) -> bool:
    return content_type in {"", "application/octet-stream"}


def _post_process_extracted_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    cleaned_lines = []
    seen = set()
    for line in lines:
        if not line:
            continue
        normalized = " ".join(line.split()).lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        cleaned_lines.append(" ".join(line.split()))
    return "\n".join(cleaned_lines).strip()


async def save_and_extract_text(upload_file: UploadFile, destination_dir: Path) -> Tuple[Path, str]:
    suffix = Path(upload_file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    content_type = (upload_file.content_type or "").lower()
    allowed_content_types = ALLOWED_MIME_TYPES_BY_EXTENSION.get(suffix, set())
    if content_type and not _is_browser_default_mime_type(content_type) and content_type not in allowed_content_types:
        logger.warning(
            "Rejected upload due to unsupported MIME type (filename=%s content_type=%s extension=%s)",
            upload_file.filename,
            content_type,
            suffix,
        )
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported content type for {suffix}: {content_type}",
        )

    destination_dir.mkdir(parents=True, exist_ok=True)
    original_name = Path(upload_file.filename or f"resume{suffix}").stem
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", original_name).strip("._") or "resume"
    try:
        content = await upload_file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")
        if len(content) > MAX_UPLOAD_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"Uploaded file is too large. Maximum allowed size is {settings.max_upload_mb} MB.",
            )

        file_hash = hashlib.sha256(content).hexdigest()[:16]
        file_path = destination_dir / f"{safe_name}-{file_hash}{suffix}"
        file_path.write_bytes(content)

        if suffix == ".pdf":
            text = extract_text_from_pdf_bytes(content)
        else:
            text = extract_text_from_docx_bytes(content)

        text = _post_process_extracted_text(text)
        if not text:
            raise HTTPException(status_code=400, detail="Could not extract readable text from the uploaded file")

        logger.info(
            "Saved resume upload to %s (filename=%s content_type=%s bytes=%s)",
            file_path,
            upload_file.filename,
            content_type,
            len(content),
        )
        return file_path, text.strip()
    finally:
        await upload_file.close()


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    pages.append(page_text)
        text = "\n".join(pages)
        logger.info(f"Extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        logger.exception("PDF extraction failed")
        raise HTTPException(status_code=400, detail=f"Failed to extract PDF: {str(e)}")


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        table_cells.append(cell.text.strip())
        text = "\n".join(paragraphs + table_cells)
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.exception("DOCX extraction failed")
        raise HTTPException(status_code=400, detail=f"Failed to extract DOCX: {str(e)}")
